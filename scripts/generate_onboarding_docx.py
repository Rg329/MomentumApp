"""Generate ONBOARDING_DATA_FLOW.docx from report content."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

OUT = Path(__file__).resolve().parent.parent / "docs" / "ONBOARDING_DATA_FLOW.docx"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Momentum App — Onboarding Data Flow Report", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    meta = doc.add_paragraph()
    meta.add_run("Project: ").bold = True
    meta.add_run("MomentumApp (Expo SDK 56 · React Native)\n")
    meta.add_run("Date: ").bold = True
    meta.add_run("June 2026\n")
    meta.add_run("Purpose: ").bold = True
    meta.add_run(
        "Complete map of what onboarding collects, where it is stored, and what reaches Supabase."
    )

    doc.add_heading("Executive summary", level=1)
    for bullet in [
        "Onboarding has 4 steps (procrastination pattern, peak focus time, circadian rhythm, coaching style).",
        "All answers are saved locally on the device via Zustand + AsyncStorage.",
        "Only 3 quiz answers (+ name/email from a later screen) are intended for Supabase profiles — and only if the user signs in and completes the Credentials screen.",
        "Wake/sleep times are never sent to Supabase (no DB columns exist).",
        "Derived personalization (greetings, coaching copy, schedule hints) is computed at runtime and never persisted.",
        "Users can skip Auth and skip Credentials, so cloud persistence is often never reached.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Navigation flow", level=1)
    flow = doc.add_paragraph()
    flow.add_run(
        "Splash → Onboarding (4 steps) → Auth → Credentials → Pro Offer → Main Tabs\n"
        "Auth can skip to Pro Offer. Credentials can skip to Pro Offer.\n\n"
        "Key files: src/navigation/RootNavigator.tsx, OnboardingScreen.tsx, AuthScreen.tsx, CredentialsScreen.tsx\n\n"
        "Note: hasOnboarded is set when onboarding finishes but is not used to skip onboarding on return visits — Splash always shows Get Started again."
    )

    doc.add_heading("1. Every onboarding question collected", level=1)
    doc.add_paragraph(
        "Defined in src/screens/OnboardingScreen.tsx and src/data/mockData.ts (ONBOARDING_OPTIONS)."
    )

    doc.add_heading("Step 1 — Procrastination pattern", level=2)
    doc.add_paragraph('Question: "What usually makes you procrastinate?"')
    doc.add_paragraph("Type: Single select")
    doc.add_paragraph("Stored as: onboardingData.procrastinationType")
    add_table(
        doc,
        ["Key", "Label shown to user"],
        [
            ("overwhelmed_tasks", "I feel overwhelmed by large tasks"),
            ("waiting_motivation", "I keep waiting for motivation"),
            ("dont_know_start", "I don't know where to start"),
            ("easily_distracted", "I get distracted easily"),
            ("changing_plans", "I keep changing my plans"),
            ("underestimate_time", "I fear I will fail"),
        ],
    )

    doc.add_heading("Step 2 — Peak focus time", level=2)
    doc.add_paragraph('Question: "When are you naturally most focused?"')
    doc.add_paragraph("Type: Single select | Stored as: onboardingData.peakTime")
    add_table(
        doc,
        ["Key", "Label"],
        [("morning", "Morning"), ("afternoon", "Afternoon"), ("evening", "Evening")],
    )

    doc.add_heading("Step 3 — Circadian rhythm", level=2)
    doc.add_paragraph('Question: "What is your circadian rhythm?"')
    doc.add_paragraph("Type: Two sliders (CircadianRhythmPicker.tsx)")
    doc.add_paragraph("Stored as: wakeTime and sleepTime (top-level store fields, not inside onboardingData)")
    add_table(
        doc,
        ["Field", "Format", "Range", "Default"],
        [
            ("Wake time", "Minutes from midnight", "5:00 AM – 10:00 AM (300–600), step 15 min", "390 (6:30 AM)"),
            ("Sleep time", "Minutes from midnight", "10:00 PM – 5:00 AM (1320–1740), step 15 min", "1365 (~10:45 PM)"),
        ],
    )
    doc.add_paragraph("Written to the store on each slider move, not batched at Finish.")

    doc.add_heading("Step 4 — Coaching style", level=2)
    doc.add_paragraph('Question: "How should Momentum coach you?"')
    doc.add_paragraph("Type: Single select (with descriptions) | Stored as: onboardingData.coaching")
    add_table(
        doc,
        ["Key", "Label"],
        [("supportive", "Supportive"), ("balanced", "Balanced"), ("strict", "Strict")],
    )

    doc.add_heading("2. Where each answer is stored", level=1)

    doc.add_heading("A. Local device (primary storage)", level=2)
    doc.add_paragraph("Mechanism: Zustand persist → AsyncStorage")
    doc.add_paragraph("Storage key: momentum-app-store")
    doc.add_paragraph("File: src/store/useAppStore.ts")
    add_table(
        doc,
        ["Data", "Zustand field", "When written"],
        [
            ("Procrastination type", "onboardingData.procrastinationType", 'Onboarding "Finish"'),
            ("Peak time", "onboardingData.peakTime", 'Onboarding "Finish"'),
            ("Coaching style", "onboardingData.coaching", 'Onboarding "Finish"'),
            ("Wake time", "wakeTime", "Each slider change (step 3)"),
            ("Sleep time", "sleepTime", "Each slider change (step 3)"),
            ("Onboarding done flag", "hasOnboarded", 'Onboarding "Finish"'),
        ],
    )

    doc.add_heading("B. In-memory only (during onboarding)", level=2)
    add_table(
        doc,
        ["Data", "Location", "Lifetime"],
        [
            ("Current step, UI selections", "OnboardingScreen local state", "Until finish"),
            ("Success overlay", "showSuccess", "Until navigate to Auth"),
        ],
    )

    doc.add_heading("C. Derived at runtime (not stored)", level=2)
    add_table(
        doc,
        ["Output", "Source inputs", "Used by"],
        [
            (
                "PersonalizationContext — greetings, coaching messages, schedule hints, badges",
                "procrastinationType, peakTime, coaching only",
                "usePersonalization() across app",
            ),
        ],
    )
    doc.add_paragraph(
        "Files: src/personalization/engine.ts, usePersonalization.ts, content.ts. "
        "Wake/sleep feed Brain Dump capacity UI only — not the personalization engine."
    )

    doc.add_heading("D. Post-onboarding (related screens)", level=2)
    add_table(
        doc,
        ["Screen", "Data collected", "Local storage", "Supabase"],
        [
            ("Auth", "Email (magic link sign-in)", "account.email", "auth.users session"),
            ("Credentials", "Name, email", "account.name, account.email, account.createdAt", "profiles upsert (best-effort)"),
        ],
    )

    doc.add_heading("3. Supabase tables involved", level=1)
    add_table(
        doc,
        ["Table", "Role"],
        [
            ("auth.users", "Created when user completes magic-link sign-in"),
            ("public.profiles", "One row per user; holds profile + onboarding quiz answers"),
        ],
    )
    doc.add_paragraph("No SQL migration files in the repo. The profiles table was created manually in the Supabase SQL Editor.")
    doc.add_paragraph("App code: src/supabase/client.ts, src/repositories/profileRepo.ts (upsertMyProfile), src/screens/CredentialsScreen.tsx (only caller).")
    doc.add_paragraph("OnboardingScreen does NOT write to Supabase on finish.")

    doc.add_heading("4. Database columns used", level=1)
    doc.add_heading("public.profiles schema (Phase 1 SQL)", level=2)
    sql = doc.add_paragraph()
    sql.add_run(
        "create table public.profiles (\n"
        "  id uuid primary key references auth.users(id) on delete cascade,\n"
        "  email text, name text,\n"
        "  procrastination_type text, peak_time text, coach_style text,\n"
        "  created_at timestamptz not null default now(),\n"
        "  updated_at timestamptz not null default now()\n"
        ");"
    ).font.name = "Consolas"

    doc.add_heading("What the app writes", level=2)
    add_table(
        doc,
        ["Column", "Set by app?", "Value source"],
        [
            ("id", "Yes", "auth.user.id"),
            ("email", "Yes", "Credentials screen / auth user"),
            ("name", "Yes", "Credentials screen"),
            ("procrastination_type", "Yes", "onboardingData.procrastinationType"),
            ("peak_time", "Yes", "onboardingData.peakTime"),
            ("coach_style", "Yes", "onboardingData.coaching"),
            ("created_at", "DB default", "Not set by app on insert"),
            ("updated_at", "Yes", "ISO timestamp on upsert"),
        ],
    )

    doc.add_heading("Row Level Security (RLS)", level=2)
    for bullet in [
        "profiles_select_own — authenticated users read own row (id = auth.uid())",
        "profiles_insert_own — authenticated users insert own row",
        "profiles_update_own — authenticated users update own row",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("5. Data that is NOT persisted (or only partially)", level=1)

    doc.add_heading("Never sent to Supabase", level=2)
    add_table(
        doc,
        ["Data", "Local only?", "Notes"],
        [
            ("Wake time", "Yes", "No wake_time column in profiles"),
            ("Sleep time", "Yes", "No sleep_time column in profiles"),
            ("Derived personalization", "Recomputed each session", "Greetings, hints, coaching copy"),
            ("hasOnboarded", "Yes", "Not in DB; not used for routing"),
        ],
    )

    doc.add_heading("Intended for Supabase but often never saved", level=2)
    add_table(
        doc,
        ["Data", "Why it may not reach the cloud"],
        [
            ("procrastination_type, peak_time, coach_style", "Upsert only on Credentials Continue; requires auth; errors silently ignored; user can skip Auth or Credentials"),
            ("name, email", "Same conditions as above"),
        ],
    )

    doc.add_heading("Other gaps", level=2)
    for bullet in [
        "User can repeat onboarding from Splash; later run overwrites local store.",
        "No sync from Supabase back to the app on login (one-way upsert only).",
        "Schedule, constraints, tasks, insights remain local or mock — out of onboarding scope.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Persistence matrix (quick reference)", level=1)
    add_table(
        doc,
        ["Field", "Collected in onboarding?", "Local (device)", "Supabase profiles"],
        [
            ("Procrastination type", "Yes", "Yes", "Only if Auth + Credentials"),
            ("Peak time", "Yes", "Yes", "Only if Auth + Credentials"),
            ("Coaching style", "Yes", "Yes", "Only if Auth + Credentials"),
            ("Wake time", "Yes", "Yes", "No"),
            ("Sleep time", "Yes", "Yes", "No"),
            ("hasOnboarded", "Yes (flag)", "Yes", "No"),
            ("Name", "No (Credentials)", "Yes", "Only if Auth + Credentials"),
            ("Email", "No (Auth/Credentials)", "Yes", "Only if Auth + Credentials"),
            ("Personalization output", "Derived", "Recomputed", "No"),
        ],
    )

    doc.add_heading("Key source files", level=1)
    add_table(
        doc,
        ["File", "Responsibility"],
        [
            ("src/screens/OnboardingScreen.tsx", "4-step UI, local state, finish handler"),
            ("src/data/mockData.ts", "Question options (ONBOARDING_OPTIONS)"),
            ("src/components/CircadianRhythmPicker.tsx", "Wake/sleep sliders"),
            ("src/store/useAppStore.ts", "Persistent local state"),
            ("src/personalization/engine.ts", "Derives UX from quiz answers"),
            ("src/screens/AuthScreen.tsx", "Magic-link sign-in (optional)"),
            ("src/screens/CredentialsScreen.tsx", "Name/email + Supabase upsert"),
            ("src/repositories/profileRepo.ts", "profiles table upsert"),
            ("src/supabase/client.ts", "Supabase client + auth sessions"),
        ],
    )

    doc.add_heading("Recommended next steps (for discussion)", level=1)
    for i, step in enumerate(
        [
            "Add wake_time and sleep_time columns to profiles and include them in upsertMyProfile.",
            "Call profile upsert after onboarding and after auth, not only on Credentials.",
            "Surface upsert errors instead of swallowing them.",
            "Use hasOnboarded (or Supabase profile fetch) to skip onboarding for returning users.",
            "Load profile from Supabase on sign-in so data survives device changes.",
        ],
        1,
    ):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_paragraph()
    doc.add_paragraph(
        "Generated from codebase analysis of MomentumApp. No code changes were made for this document."
    ).italic = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
